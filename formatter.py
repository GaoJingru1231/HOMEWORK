"""格式快调 - 核心排版逻辑（仅依赖 python-docx + 标准库）。

供两处复用：
  - 本地/同源 FastAPI 服务 ``server.py``
  - 腾讯云 SCF 云函数 ``cloudfunctions/format_api/index.py``

刻意不依赖 fastapi / pydantic，以便 SCF 函数包只需打包 python-docx（及 lxml）。
"""

import io
import re
from copy import deepcopy
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ===== 数据模型（纯 Python 类，避免 pydantic 依赖）=====

class FormatParams:
    """格式参数容器。"""

    __fields__ = (
        "font_cn", "font_en", "font_size", "line_spacing", "first_line_indent",
        "alignment", "space_before", "space_after",
        "margin_top", "margin_bottom", "margin_left", "margin_right",
        "h1_size", "h1_bold", "h2_size", "h2_bold",
    )

    def __init__(self, **kwargs):
        # 默认值（与原 pydantic 模型一致）
        self.font_cn = "宋体"
        self.font_en = "Times New Roman"
        self.font_size = 12.0
        self.line_spacing = 1.5
        self.first_line_indent = 2
        self.alignment = "JUSTIFY"
        self.space_before = 0.0
        self.space_after = 0.0
        self.margin_top = 2.54
        self.margin_bottom = 2.54
        self.margin_left = 3.17
        self.margin_right = 3.17
        self.h1_size = 16.0
        self.h1_bold = True
        self.h2_size = 14.0
        self.h2_bold = True
        for k, v in kwargs.items():
            setattr(self, k, v)

    def to_dict(self):
        return {f: getattr(self, f) for f in self.__fields__}

    def copy(self):
        return deepcopy(self)


class DetectionIssue:
    def __init__(self, type, category, message, count, details=None):
        self.type = type  # "fixed" | "warning"
        self.category = category
        self.message = message
        self.count = count
        self.details = details or []

    def to_dict(self):
        return {
            "type": self.type,
            "category": self.category,
            "message": self.message,
            "count": self.count,
            "details": self.details,
        }


class DetectionReport:
    def __init__(self, total_issues, fixed_count, warning_count, issues):
        self.total_issues = total_issues
        self.fixed_count = fixed_count
        self.warning_count = warning_count
        self.issues = issues

    def to_dict(self):
        return {
            "total_issues": self.total_issues,
            "fixed_count": self.fixed_count,
            "warning_count": self.warning_count,
            "issues": [i.to_dict() if hasattr(i, "to_dict") else i for i in self.issues],
        }


# ===== 预设模板（12 套，与前端 builtinTemplates 一一对应）=====
# 字号说明：小四=12pt，五号=10.5pt，小二=18pt，三号=16pt，四号=14pt

PRESET_TEMPLATES = {
    # t1 本科毕业论文
    "undergrad_thesis": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t2 硕士毕业论文
    "master_thesis": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=3.0, margin_bottom=2.5, margin_left=3.0, margin_right=3.0,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t3 课程论文（文科）
    "course_paper_arts": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.25, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t4 课程论文（理工科）
    "course_paper_sci": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.0, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t5 实验报告
    "lab_report": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.25, first_line_indent=0, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=2.54, margin_right=2.54,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t6 开题报告
    "proposal_report": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.0, margin_right=3.0,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t7 学术期刊投稿
    "journal_submission": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.25, first_line_indent=0, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=2.54, margin_right=2.54,
        h1_size=12.0, h1_bold=True, h2_size=10.5, h2_bold=True,
    ),
    # t8 小组作业报告
    "group_report": FormatParams(
        font_cn="微软雅黑", font_en="Calibri", font_size=12.0,
        line_spacing=1.25, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t9 读书笔记
    "reading_notes": FormatParams(
        font_cn="楷体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="LEFT",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=False,
    ),
    # t10 学年论文
    "annual_paper": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t11 文献综述
    "literature_review": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.25, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=12.0, h1_bold=True, h2_size=10.5, h2_bold=True,
    ),
    # t12 答辩PPT提纲
    "defense_outline": FormatParams(
        font_cn="黑体", font_en="Arial", font_size=12.0,
        line_spacing=1.25, first_line_indent=0, alignment="LEFT",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
}


# ===== 辅助函数 =====

ALIGNMENT_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def set_cell_font(run, font_cn: str, font_en: str, size_pt: float):
    """设置 run 的中英文字体、字号"""
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def set_paragraph_format(paragraph, params: FormatParams):
    """设置段落格式（行距/段前段后/对齐）。首行缩进由 set_first_line_indent_chars 处理。"""
    pf = paragraph.paragraph_format
    pf.line_spacing = params.line_spacing
    pf.space_before = Pt(params.space_before)
    pf.space_after = Pt(params.space_after)
    pf.alignment = ALIGNMENT_MAP.get(params.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)


def set_first_line_indent_chars(paragraph, chars: int, font_size_pt: float):
    """设置首行缩进（字符单位，匹配 Word「N字符」）。
    chars>0: firstLineChars=N*100 + firstLine(twips) 兜底；
    chars==0: 显式清零，覆盖样式继承与残留字符缩进。"""
    pf = paragraph.paragraph_format
    if chars and chars > 0:
        pf.first_line_indent = Pt(font_size_pt * chars)  # 触发 python-docx 按 schema 创建 w:ind
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        pf.first_line_indent = Pt(0)
        ind = pPr.find(qn('w:ind'))
    if ind is None:
        return
    for attr in ('firstLine', 'hanging', 'firstLineChars', 'hangingChars'):
        key = qn(f'w:{attr}')
        if key in ind.attrib:
            del ind.attrib[key]
    if chars and chars > 0:
        ind.set(qn('w:firstLineChars'), str(int(chars * 100)))
        ind.set(qn('w:firstLine'), str(int(font_size_pt * chars * 20)))  # 1pt=20twips
    else:
        ind.set(qn('w:firstLine'), '0')


# 智能标题检测正则
HEADING_PATTERNS = [
    re.compile(r'^(第[一二三四五六七八九十百千\d]+章|第[一二三四五六七八九十百千\d]+节)\s'),
    re.compile(r'^[一二三四五六七八九十]+[、，.\s]'),
    re.compile(r'^[\d]+[\.\、\s]+'),  # 1. 1、1
    re.compile(r'^(摘要|Abstract|关键词|Keywords|引言|绪论|前言|结论|总结|展望|致谢|参考文献|附录|附录[一二三四五六七八九十\d]+)'),
]


def is_heading_paragraph(text: str, bold: bool) -> int:
    """检测段落是否为非标准标题，返回标题层级(1/2/3)，0表示不是标题"""
    text = text.strip()
    if not text:
        return 0
    # 短段落 + 加粗 → 大概率是标题
    if len(text) < 40 and bold:
        for pattern in HEADING_PATTERNS:
            if pattern.search(text):
                return 1
        return 2  # 短加粗段，可能是二级标题
    # 匹配标题模式
    for pattern in HEADING_PATTERNS:
        if pattern.search(text):
            return 1 if len(text) < 40 else 2
    return 0


def apply_document_format(doc: Document, params: FormatParams) -> int:
    """对文档全文应用格式参数，返回修改段落数"""
    modified_count = 0

    # 更新 Normal 默认样式字体，让无 run 段落/空段继承模板字体
    try:
        normal = doc.styles['Normal']
        normal.font.size = Pt(params.font_size)
        normal.font.name = params.font_en
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:eastAsia'), params.font_cn)
        rfonts.set(qn('w:ascii'), params.font_en)
        rfonts.set(qn('w:hAnsi'), params.font_en)
    except Exception:
        pass

    heading_styles = {
        'Heading 1': (Pt(params.h1_size), params.h1_bold),
        'Heading 2': (Pt(params.h2_size), params.h2_bold),
        'Heading 3': (Pt(params.font_size), True),
    }

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()

        if style_name in heading_styles:
            size, bold = heading_styles[style_name]
            for run in paragraph.runs:
                run.font.size = size
                run.bold = bold
                set_cell_font(run, "黑体", params.font_en, size.pt)
            set_first_line_indent_chars(paragraph, 0, params.font_size)
            modified_count += 1
        elif style_name != 'Heading 1' and style_name != 'Heading 2':
            # 检查段落是否已有加粗 run
            has_bold = any(run.bold for run in paragraph.runs if run.bold)
            heading_level = is_heading_paragraph(text, has_bold)

            if heading_level == 1:
                # 按一级标题处理
                for run in paragraph.runs:
                    run.font.size = Pt(params.h1_size)
                    run.bold = params.h1_bold
                    set_cell_font(run, "黑体", params.font_en, params.h1_size)
                # 标题不需要首行缩进
                set_first_line_indent_chars(paragraph, 0, params.font_size)
                paragraph.paragraph_format.line_spacing = params.line_spacing
                modified_count += 1
            elif heading_level == 2:
                for run in paragraph.runs:
                    run.font.size = Pt(params.h2_size)
                    run.bold = params.h2_bold
                    set_cell_font(run, "黑体", params.font_en, params.h2_size)
                set_first_line_indent_chars(paragraph, 0, params.font_size)
                paragraph.paragraph_format.line_spacing = params.line_spacing
                modified_count += 1
            else:
                for run in paragraph.runs:
                    set_cell_font(run, params.font_cn, params.font_en, params.font_size)
                set_paragraph_format(paragraph, params)
                set_first_line_indent_chars(paragraph, params.first_line_indent, params.font_size)
                modified_count += 1

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(params.margin_top)
        section.bottom_margin = Cm(params.margin_bottom)
        section.left_margin = Cm(params.margin_left)
        section.right_margin = Cm(params.margin_right)

    # 表格内文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_cell_font(run, params.font_cn, params.font_en, params.font_size)
                    set_paragraph_format(paragraph, params)
                    set_first_line_indent_chars(paragraph, params.first_line_indent, params.font_size)

    return modified_count


def detect_format_issues(doc: Document, params: FormatParams) -> DetectionReport:
    """检测文档中的格式问题"""
    issues: list[DetectionIssue] = []

    # 字体检测
    font_set = set()
    font_inconsistencies = []
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            if run.font.name:
                font_set.add(run.font.name)
    if len(font_set) > 1:
        font_inconsistencies = sorted(font_set)
    if font_inconsistencies:
        issues.append(DetectionIssue(
            type="fixed", category="font",
            message=f"检测到 {len(font_set)} 种字体，建议统一为「{params.font_cn}」",
            count=len(font_set), details=font_inconsistencies
        ))

    # 行距检测
    spacing_values = set()
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.line_spacing:
            spacing_values.add(round(pf.line_spacing, 2))
    if len(spacing_values) > 1:
        issues.append(DetectionIssue(
            type="fixed", category="line_spacing",
            message=f"行距不一致，存在 {len(spacing_values)} 种不同行距值",
            count=len(spacing_values),
            details=[str(v) for v in sorted(spacing_values)]
        ))

    # 缩进检测
    no_indent_count = 0
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.first_line_indent is None and p.text.strip():
            no_indent_count += 1
    if no_indent_count > 0:
        issues.append(DetectionIssue(
            type="fixed", category="indent",
            message=f"检测到 {no_indent_count} 段缺少首行缩进",
            count=no_indent_count,
            details=[f"将统一设置为首行缩进 {params.first_line_indent} 字符"]
        ))

    # 标题层级检测
    heading_styles_found = set()
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith('Heading'):
            heading_styles_found.add(p.style.name)
    if heading_styles_found:
        issues.append(DetectionIssue(
            type="fixed", category="headings",
            message=f"检测到 {len(heading_styles_found)} 级标题样式，将按模板统一格式",
            count=len(heading_styles_found),
            details=sorted(heading_styles_found)
        ))

    # 页边距检测
    if doc.sections:
        sec = doc.sections[0]
        margin_issues = []
        if abs(sec.top_margin.cm - params.margin_top) > 0.2:
            margin_issues.append(f"上边距 {sec.top_margin.cm:.1f}cm → {params.margin_top}cm")
        if abs(sec.left_margin.cm - params.margin_left) > 0.2:
            margin_issues.append(f"左边距 {sec.left_margin.cm:.1f}cm → {params.margin_left}cm")
        if margin_issues:
            issues.append(DetectionIssue(
                type="fixed", category="margin",
                message="页边距不符合规范，将自动修正",
                count=len(margin_issues), details=margin_issues
            ))

    # 空行检测
    empty_count = sum(1 for p in doc.paragraphs if not p.text.strip() and p.paragraph_format.space_before is None)
    if empty_count > 5:
        issues.append(DetectionIssue(
            type="warning", category="empty_lines",
            message=f"检测到 {empty_count} 处多余空行/空白段落",
            count=empty_count,
            details=["建议手动检查并删除不必要的空行"]
        ))

    fixed = sum(1 for i in issues if i.type == "fixed")
    warning = sum(1 for i in issues if i.type == "warning")

    return DetectionReport(
        total_issues=len(issues),
        fixed_count=fixed,
        warning_count=warning,
        issues=issues
    )


# ===== 文件格式验证 =====

ACCEPTED_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template",
    "application/msword",
    "application/x-msword",
    "application/octet-stream",
    "application/zip",
    "application/x-zip-compressed",
    "binary/octet-stream",
    "",
    None,
}


def is_valid_doc_file(filename: Optional[str], content_type: Optional[str]) -> bool:
    """判断是否为可接受的 Word 文档格式"""
    if filename:
        lower_name = filename.lower()
        if lower_name.endswith('.docx') or lower_name.endswith('.doc'):
            return True
    if content_type in ACCEPTED_MIME_TYPES:
        return True
    if content_type:
        ct_base = content_type.split(';')[0].strip().lower()
        if ct_base in ACCEPTED_MIME_TYPES:
            return True
    return False


def resolve_params(template: str, overrides: dict) -> FormatParams:
    """根据模板 key + 可选 override 字典构造 FormatParams。
    overrides 中值为 None / 空字符串的字段会被忽略。"""
    params = (PRESET_TEMPLATES.get(template) or PRESET_TEMPLATES["undergrad_thesis"]).copy()
    for field in FormatParams.__fields__:
        v = overrides.get(field)
        if v is None or v == "":
            continue
        # 类型转换
        if field in ("font_size", "line_spacing", "space_before", "space_after",
                     "margin_top", "margin_bottom", "margin_left", "margin_right",
                     "h1_size", "h2_size"):
            try:
                v = float(v)
            except (TypeError, ValueError):
                continue
        elif field in ("first_line_indent",):
            try:
                v = int(float(v))
            except (TypeError, ValueError):
                continue
        elif field in ("h1_bold", "h2_bold"):
            if isinstance(v, str):
                v = v.strip().lower() in ("true", "1", "yes", "on")
        setattr(params, field, v)
    return params
