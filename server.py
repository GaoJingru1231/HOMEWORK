"""
格式快调 - 后端 API 服务
基于 FastAPI + python-docx 实现文档格式检测与一键排版
"""

import io
import os
import uuid
import shutil
from copy import deepcopy
from pathlib import Path
from datetime import datetime
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# python-docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

app = FastAPI(title="格式快调 API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 配置
BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
PROCESSED_DIR = BASE_DIR / "processed"
UPLOAD_DIR.mkdir(exist_ok=True)
PROCESSED_DIR.mkdir(exist_ok=True)

# ===== 数据模型 =====

class FormatParams(BaseModel):
    font_cn: str = "宋体"
    font_en: str = "Times New Roman"
    font_size: float = 12.0
    line_spacing: float = 1.5
    first_line_indent: int = 2
    alignment: str = "JUSTIFY"
    space_before: float = 0.0
    space_after: float = 0.0
    margin_top: float = 2.54
    margin_bottom: float = 2.54
    margin_left: float = 3.17
    margin_right: float = 3.17
    h1_size: float = 16.0
    h1_bold: bool = True
    h2_size: float = 14.0
    h2_bold: bool = True

class DetectionIssue(BaseModel):
    type: str  # "fixed" | "warning"
    category: str
    message: str
    count: int
    details: list[str] = []

class DetectionReport(BaseModel):
    total_issues: int
    fixed_count: int
    warning_count: int
    issues: list[DetectionIssue]

# ===== 预设模板（12套，与前端 builtinTemplates 一一对应）=====
# 字号说明：小四=12pt，五号=10.5pt，小二=18pt，三号=16pt，四号=14pt

PRESET_TEMPLATES = {
    # t1 本科毕业论文
    "undergrad_thesis": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t2 硕士毕业论文
    "master_thesis": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=3.0, margin_bottom=2.5, margin_left=3.0, margin_right=3.0,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t3 课程论文（文科）
    "course_paper_arts": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.25, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t4 课程论文（理工科）
    "course_paper_sci": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.0, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t5 实验报告
    "lab_report": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.25, first_line_indent=0, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=2.54, margin_right=2.54,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t6 开题报告
    "proposal_report": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.0, margin_right=3.0,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t7 学术期刊投稿
    "journal_submission": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.25, first_line_indent=0, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=2.54, margin_right=2.54,
        h1_size=12.0, h1_bold=True, h2_size=10.5, h2_bold=True,
    ),
    # t8 小组作业报告
    "group_report": FormatParams(
        font_cn="微软雅黑", font_en="Calibri", font_size=12.0,
        line_spacing=1.25, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
    # t9 读书笔记
    "reading_notes": FormatParams(
        font_cn="楷体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="LEFT",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=False,
    ),
    # t10 学年论文
    "annual_paper": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=12.0,
        line_spacing=1.5, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=16.0, h1_bold=True, h2_size=14.0, h2_bold=True,
    ),
    # t11 文献综述
    "literature_review": FormatParams(
        font_cn="宋体", font_en="Times New Roman", font_size=10.5,
        line_spacing=1.25, first_line_indent=2, alignment="JUSTIFY",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=12.0, h1_bold=True, h2_size=10.5, h2_bold=True,
    ),
    # t12 答辩PPT提纲
    "defense_outline": FormatParams(
        font_cn="黑体", font_en="Arial", font_size=12.0,
        line_spacing=1.25, first_line_indent=0, alignment="LEFT",
        space_before=0, space_after=0,
        margin_top=2.54, margin_bottom=2.54, margin_left=3.17, margin_right=3.17,
        h1_size=14.0, h1_bold=True, h2_size=12.0, h2_bold=True,
    ),
}

# ===== 辅助函数 =====

ALIGNMENT_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def set_cell_font(run, font_cn: str, font_en: str, size_pt: float):
    """设置 run 的中英文字体、字号"""
    run.font.size = Pt(size_pt)
    run.font.name = font_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def set_paragraph_format(paragraph, params: FormatParams):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.line_spacing = params.line_spacing
    pf.space_before = Pt(params.space_before)
    pf.space_after = Pt(params.space_after)
    pf.alignment = ALIGNMENT_MAP.get(params.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if params.first_line_indent > 0:
        pf.first_line_indent = Pt(params.font_size * params.first_line_indent)


import re

# 智能标题检测正则
HEADING_PATTERNS = [
    re.compile(r'^(第[一二三四五六七八九十百千\d]+章|第[一二三四五六七八九十百千\d]+节)\s'),
    re.compile(r'^[一二三四五六七八九十]+[、，.\s]'),
    re.compile(r'^[\d]+[\.\、\s]+'),  # 1. 1、1 
    re.compile(r'^(摘要|Abstract|关键词|Keywords|引言|绪论|前言|结论|总结|展望|致谢|参考文献|附录|附录[一二三四五六七八九十\d]+)'),
]

def is_heading_paragraph(text: str, bold: bool) -> int:
    """检测段落是否为非标准标题，返回标题层级(1/2/3)，0表示不是标题"""
    text = text.strip()
    if not text:
        return 0
    # 短段落 + 加粗 → 大概率是标题
    if len(text) < 40 and bold:
        for pattern in HEADING_PATTERNS:
            if pattern.search(text):
                return 1
        return 2  # 短加粗段，可能是二级标题
    # 匹配标题模式
    for pattern in HEADING_PATTERNS:
        if pattern.search(text):
            return 1 if len(text) < 40 else 2
    return 0


def apply_document_format(doc: Document, params: FormatParams) -> int:
    """对文档全文应用格式参数，返回修改段落数"""
    modified_count = 0
    heading_styles = {
        'Heading 1': (Pt(params.h1_size), params.h1_bold),
        'Heading 2': (Pt(params.h2_size), params.h2_bold),
        'Heading 3': (Pt(params.font_size), True),
    }

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()

        if style_name in heading_styles:
            size, bold = heading_styles[style_name]
            for run in paragraph.runs:
                run.font.size = size
                run.bold = bold
                set_cell_font(run, "黑体", params.font_en, size.pt)
            modified_count += 1
        elif style_name != 'Heading 1' and style_name != 'Heading 2':
            # 检查段落是否已有加粗 run
            has_bold = any(run.bold for run in paragraph.runs if run.bold)
            heading_level = is_heading_paragraph(text, has_bold)

            if heading_level == 1:
                # 按一级标题处理
                for run in paragraph.runs:
                    run.font.size = Pt(params.h1_size)
                    run.bold = params.h1_bold
                    set_cell_font(run, "黑体", params.font_en, params.h1_size)
                # 标题不需要首行缩进
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.line_spacing = params.line_spacing
                modified_count += 1
            elif heading_level == 2:
                for run in paragraph.runs:
                    run.font.size = Pt(params.h2_size)
                    run.bold = params.h2_bold
                    set_cell_font(run, "黑体", params.font_en, params.h2_size)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.line_spacing = params.line_spacing
                modified_count += 1
            else:
                for run in paragraph.runs:
                    set_cell_font(run, params.font_cn, params.font_en, params.font_size)
                set_paragraph_format(paragraph, params)
                modified_count += 1

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(params.margin_top)
        section.bottom_margin = Cm(params.margin_bottom)
        section.left_margin = Cm(params.margin_left)
        section.right_margin = Cm(params.margin_right)

    # 表格内文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_cell_font(run, params.font_cn, params.font_en, params.font_size)
                    set_paragraph_format(paragraph, params)

    return modified_count


def detect_format_issues(doc: Document, params: FormatParams) -> DetectionReport:
    """检测文档中的格式问题"""
    issues: list[DetectionIssue] = []

    # 字体检测
    font_set = set()
    font_inconsistencies = []
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            if run.font.name:
                font_set.add(run.font.name)
    if len(font_set) > 1:
        font_inconsistencies = sorted(font_set)
    if font_inconsistencies:
        issues.append(DetectionIssue(
            type="fixed", category="font",
            message=f"检测到 {len(font_set)} 种字体，建议统一为「{params.font_cn}」",
            count=len(font_set), details=font_inconsistencies
        ))

    # 行距检测
    spacing_values = set()
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.line_spacing:
            spacing_values.add(round(pf.line_spacing, 2))
    if len(spacing_values) > 1:
        issues.append(DetectionIssue(
            type="fixed", category="line_spacing",
            message=f"行距不一致，存在 {len(spacing_values)} 种不同行距值",
            count=len(spacing_values),
            details=[str(v) for v in sorted(spacing_values)]
        ))

    # 缩进检测
    no_indent_count = 0
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.first_line_indent is None and p.text.strip():
            no_indent_count += 1
    if no_indent_count > 0:
        issues.append(DetectionIssue(
            type="fixed", category="indent",
            message=f"检测到 {no_indent_count} 段缺少首行缩进",
            count=no_indent_count,
            details=[f"将统一设置为首行缩进 {params.first_line_indent} 字符"]
        ))

    # 标题层级检测
    heading_styles_found = set()
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith('Heading'):
            heading_styles_found.add(p.style.name)
    if heading_styles_found:
        issues.append(DetectionIssue(
            type="fixed", category="headings",
            message=f"检测到 {len(heading_styles_found)} 级标题样式，将按模板统一格式",
            count=len(heading_styles_found),
            details=sorted(heading_styles_found)
        ))

    # 页边距检测
    if doc.sections:
        sec = doc.sections[0]
        margin_issues = []
        if abs(sec.top_margin.cm - params.margin_top) > 0.2:
            margin_issues.append(f"上边距 {sec.top_margin.cm:.1f}cm → {params.margin_top}cm")
        if abs(sec.left_margin.cm - params.margin_left) > 0.2:
            margin_issues.append(f"左边距 {sec.left_margin.cm:.1f}cm → {params.margin_left}cm")
        if margin_issues:
            issues.append(DetectionIssue(
                type="fixed", category="margin",
                message="页边距不符合规范，将自动修正",
                count=len(margin_issues), details=margin_issues
            ))

    # 空行检测
    empty_count = sum(1 for p in doc.paragraphs if not p.text.strip() and p.paragraph_format.space_before is None)
    if empty_count > 5:
        issues.append(DetectionIssue(
            type="warning", category="empty_lines",
            message=f"检测到 {empty_count} 处多余空行/空白段落",
            count=empty_count,
            details=["建议手动检查并删除不必要的空行"]
        ))

    fixed = sum(1 for i in issues if i.type == "fixed")
    warning = sum(1 for i in issues if i.type == "warning")

    return DetectionReport(
        total_issues=len(issues),
        fixed_count=fixed,
        warning_count=warning,
        issues=issues
    )


# ===== 辅助：文件格式验证 =====

ACCEPTED_MIME_TYPES = {
    # .docx
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template",
    # .doc
    "application/msword",
    "application/x-msword",
    # 通用（某些系统/浏览器上传时返回这些）
    "application/octet-stream",
    "application/zip",
    "application/x-zip-compressed",
    "binary/octet-stream",
    "",  # 某些环境 content_type 为空
    None,
}

def is_valid_doc_file(filename: Optional[str], content_type: Optional[str]) -> bool:
    """判断是否为可接受的 Word 文档格式"""
    # 优先通过文件名后缀判断（大小写不敏感）
    if filename:
        lower_name = filename.lower()
        if lower_name.endswith('.docx') or lower_name.endswith('.doc'):
            return True
    # 通过 content_type 判断
    if content_type in ACCEPTED_MIME_TYPES:
        return True
    # content_type 可能是多值（如 "application/zip; charset=binary"）
    if content_type:
        ct_base = content_type.split(';')[0].strip().lower()
        if ct_base in ACCEPTED_MIME_TYPES:
            return True
    return False


# ===== API 路由 =====

@app.get("/api/health")
async def health_check():
    return {"status": "ok", "timestamp": datetime.now().isoformat()}


@app.get("/api/templates")
async def list_templates():
    """列出所有预设模板"""
    templates = []
    for key, params in PRESET_TEMPLATES.items():
        templates.append({
            "id": key,
            "name": key,
            "params": params.model_dump(),
        })
    return {"templates": templates}


@app.post("/api/detect")
async def detect_format(file: UploadFile = File(...)):
    """
    上传文档并检测格式问题，返回检测报告。
    使用默认"本科毕业论文"模板参数作为检测基准。
    """
    if not is_valid_doc_file(file.filename, file.content_type):
        raise HTTPException(status_code=400, detail="仅支持 .docx / .doc 格式，请确认文件后缀正确")

    contents = await file.read()
    # 二次兜底：直接尝试打开，python-docx 会抛出异常告知是否有效
    try:
        doc = Document(io.BytesIO(contents))
    except Exception:
        raise HTTPException(status_code=400, detail="文件格式无效，请上传有效的 .docx 文档")

    params = PRESET_TEMPLATES["undergrad_thesis"]
    report = detect_format_issues(doc, params)

    return {"filename": file.filename, "report": report.model_dump()}


@app.post("/api/format")
async def format_document(
    file: UploadFile = File(...),
    template: str = Form("undergrad_thesis"),
    font_cn: Optional[str] = Form(None),
    font_en: Optional[str] = Form(None),
    font_size: Optional[float] = Form(None),
    line_spacing: Optional[float] = Form(None),
    first_line_indent: Optional[int] = Form(None),
    alignment: Optional[str] = Form(None),
    space_before: Optional[float] = Form(None),
    space_after: Optional[float] = Form(None),
    margin_top: Optional[float] = Form(None),
    margin_bottom: Optional[float] = Form(None),
    margin_left: Optional[float] = Form(None),
    margin_right: Optional[float] = Form(None),
    h1_size: Optional[float] = Form(None),
    h1_bold: Optional[bool] = Form(None),
    h2_size: Optional[float] = Form(None),
    h2_bold: Optional[bool] = Form(None),
):
    """
    上传文档并按指定模板 + 可选自定义参数进行排版，返回处理后的文档。
    """
    if not is_valid_doc_file(file.filename, file.content_type):
        raise HTTPException(status_code=400, detail="仅支持 .docx / .doc 格式，请确认文件后缀正确")

    # 读取文档（二次兜底：python-docx 会报错如果文件无效）
    contents = await file.read()
    try:
        doc = Document(io.BytesIO(contents))
    except Exception:
        raise HTTPException(status_code=400, detail="文件格式无效，请上传有效的 .docx 文档")
    if template in PRESET_TEMPLATES:
        params = deepcopy(PRESET_TEMPLATES[template])
    else:
        params = deepcopy(PRESET_TEMPLATES["undergrad_thesis"])

    # 覆盖自定义参数
    override_fields = {
        "font_cn": font_cn, "font_en": font_en, "font_size": font_size,
        "line_spacing": line_spacing, "first_line_indent": first_line_indent,
        "alignment": alignment, "space_before": space_before, "space_after": space_after,
        "margin_top": margin_top, "margin_bottom": margin_bottom,
        "margin_left": margin_left, "margin_right": margin_right,
        "h1_size": h1_size, "h1_bold": h1_bold, "h2_size": h2_size, "h2_bold": h2_bold,
    }
    for field, value in override_fields.items():
        if value is not None:
            setattr(params, field, value)

    # 检测
    report = detect_format_issues(doc, params)

    # 应用格式
    modified_count = apply_document_format(doc, params)

    # 保存
    task_id = uuid.uuid4().hex[:12]
    original_name = Path(file.filename).stem
    output_name = f"{original_name}_已排版_{template}.docx"
    output_path = PROCESSED_DIR / f"{task_id}_{output_name}"
    doc.save(str(output_path))

    return JSONResponse({
        "task_id": task_id,
        "original_filename": file.filename,
        "output_filename": output_name,
        "modified_paragraphs": modified_count,
        "report": report.model_dump(),
        "download_url": f"/api/download/{task_id}/{output_name}",
        "template_used": template,
    })


@app.get("/api/download/{task_id}/{filename}")
async def download_file(task_id: str, filename: str):
    """下载处理后的文档"""
    file_path = PROCESSED_DIR / f"{task_id}_{filename}"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ===== 静态文件服务 =====
from fastapi.responses import HTMLResponse

INDEX_PATH = BASE_DIR / "index.html"

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    if INDEX_PATH.exists():
        return HTMLResponse(
            content=INDEX_PATH.read_text(encoding="utf-8"),
            headers={"Cache-Control": "no-cache, no-store, must-revalidate", "Pragma": "no-cache", "Expires": "0"}
        )
    raise HTTPException(status_code=404, detail="index.html not found")


# ===== 启动入口 =====
if __name__ == "__main__":
    import uvicorn
    import os
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("server:app", host="0.0.0.0", port=port, reload=False)
